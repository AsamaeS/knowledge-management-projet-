import io
import json
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
from uuid import UUID
from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

# Models
from backend.models.document import Document
from backend.models.chunk import DocumentChunk
from backend.models.graph import Node, Edge

# Services
from backend.core.ai_service import AIService
from backend.schemas.document import IngestionResponse

logger = logging.getLogger(__name__)

class IngestionService:
    def __init__(self, db: AsyncSession, ai_service: AIService):
        self.db = db
        self.ai_service = ai_service

    async def ingest_file(
        self,
        file: UploadFile,
        source_type: str,
        author: Optional[str] = None,
        doc_date: Optional[str] = None,
        metadata_extra: Optional[Dict[str, Any]] = None
    ) -> IngestionResponse:
        """Parses, chunks, embeds, extracts entities, and stores a document."""
        filename = file.filename
        content_bytes = await file.read()
        
        logger.info(f"Ingesting file '{filename}' of type {source_type}")
        
        # 1. Parse File Content
        raw_text = self._parse_file(filename, content_bytes)
        
        # Parse document date if provided
        parsed_date = None
        if doc_date:
            try:
                parsed_date = datetime.strptime(doc_date, "%Y-%m-%d").date()
            except ValueError:
                logger.warning(f"Failed to parse document date '{doc_date}'. Saving as None.")

        # Save metadata dict
        metadata = {"filename": filename, "source_type": source_type}
        if metadata_extra:
            metadata.update(metadata_extra)

        # 2. Create and commit Document record
        document = Document(
            filename=filename,
            source_type=source_type,
            author=author,
            doc_date=parsed_date,
            raw_text=raw_text,
            metadata=metadata
        )
        self.db.add(document)
        await self.db.flush()  # Obtain document.id
        doc_id = document.id

        try:
            # 3. Chunk the document text
            chunks = self._chunk_text(raw_text)
            chunk_count = len(chunks)
            logger.info(f"Split document into {chunk_count} chunks.")

            # 4. Generate Embeddings for Chunks
            chunk_contents = [c["content"] for c in chunks]
            embeddings = await self.ai_service.embed_batch(chunk_contents)

            # 5. Store Chunks in DB
            db_chunks = []
            for i, chunk_data in enumerate(chunks):
                db_chunk = DocumentChunk(
                    document_id=doc_id,
                    chunk_index=i,
                    content=chunk_data["content"],
                    embedding=embeddings[i],
                    token_count=chunk_data["token_count"]
                )
                self.db.add(db_chunk)
                db_chunks.append(db_chunk)
            
            await self.db.flush()

            # 6. Extract Entities and Relationships (using LLM on the whole raw text or summary)
            # If the text is very long, extract from the first few key chunks or summarize first
            # For MVP, we can run extraction on a truncated version of the text (up to 8000 tokens)
            truncated_text = raw_text[:30000] # roughly 6000-8000 tokens
            extraction_result = await self.ai_service.extract_entities(truncated_text)
            
            entities = extraction_result.get("entities", [])
            relationships = extraction_result.get("relationships", [])

            # 7. Deduplicate and Save Entities (Nodes)
            nodes_created = 0
            node_label_to_id = {}
            
            for ent in entities:
                label = ent["label"].strip()
                ent_type = ent.get("type", "concept").strip().lower()
                desc = ent.get("description", "")
                
                # Check if node already exists in DB (case-insensitive deduplication)
                existing_node_query = await self.db.execute(
                    select(Node).where(Node.label.ilike(label))
                )
                existing_node = existing_node_query.scalars().first()

                if existing_node:
                    # Update properties and add doc_id to source_ids if not present
                    if doc_id not in existing_node.source_ids:
                        existing_node.source_ids = list(existing_node.source_ids) + [doc_id]
                    # Update description if empty
                    if not existing_node.description and desc:
                        existing_node.description = desc
                    
                    node_label_to_id[label.lower()] = existing_node.id
                else:
                    # Create new node
                    node_embedding = await self.ai_service.embed(label)
                    new_node = Node(
                        label=label,
                        type=ent_type,
                        description=desc,
                        source_ids=[doc_id],
                        embedding=node_embedding,
                        properties={}
                    )
                    self.db.add(new_node)
                    await self.db.flush()  # Obtain node id
                    node_label_to_id[label.lower()] = new_node.id
                    nodes_created += 1

            # 8. Save Relationships (Edges)
            edges_created = 0
            for rel in relationships:
                src_label = rel["source"].strip().lower()
                tgt_label = rel["target"].strip().lower()
                relation = rel.get("relation", "related_to").strip().lower()
                weight = float(rel.get("weight", 1.0))
                props = rel.get("properties", {})

                # Only create edge if both nodes exist or were just created
                src_id = node_label_to_id.get(src_label)
                tgt_id = node_label_to_id.get(tgt_label)

                if src_id and tgt_id:
                    # Check if matching edge already exists
                    existing_edge_query = await self.db.execute(
                        select(Edge).where(
                            Edge.source_node == src_id,
                            Edge.target_node == tgt_id,
                            Edge.relation == relation
                        )
                    )
                    existing_edge = existing_edge_query.scalars().first()

                    if existing_edge:
                        if doc_id not in existing_edge.source_ids:
                            existing_edge.source_ids = list(existing_edge.source_ids) + [doc_id]
                        existing_edge.weight = max(existing_edge.weight, weight)
                    else:
                        new_edge = Edge(
                            source_node=src_id,
                            target_node=tgt_id,
                            relation=relation,
                            weight=weight,
                            source_ids=[doc_id],
                            properties=props
                        )
                        self.db.add(new_edge)
                        edges_created += 1

            # Complete transaction
            await self.db.commit()
            
            return IngestionResponse(
                document_id=doc_id,
                chunk_count=chunk_count,
                entities_extracted=nodes_created,
                edges_created=edges_created
            )

        except Exception as e:
            logger.error(f"Failed to ingest file: {e}")
            await self.db.rollback()
            raise

    def _parse_file(self, filename: str, content: bytes) -> str:
        """Parses binary data depending on document type."""
        ext = filename.split(".")[-1].lower()
        
        if ext == "pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(content))
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                return text
            except Exception as e:
                logger.error(f"PDF extraction error: {e}")
                raise ValueError(f"Failed to parse PDF file: {e}")
                
        elif ext == "docx":
            try:
                import docx
                doc = docx.Document(io.BytesIO(content))
                text = "\n".join([para.text for para in doc.paragraphs])
                return text
            except Exception as e:
                logger.error(f"DOCX extraction error: {e}")
                raise ValueError(f"Failed to parse DOCX file: {e}")
                
        elif ext == "txt":
            return content.decode("utf-8", errors="ignore")
            
        elif ext == "json":
            try:
                data = json.loads(content.decode("utf-8", errors="ignore"))
                return json.dumps(data, indent=2)
            except Exception as e:
                raise ValueError(f"Failed to parse JSON file: {e}")
                
        elif ext == "csv":
            return content.decode("utf-8", errors="ignore")
            
        else:
            raise ValueError(f"Unsupported file extension: .{ext}")

    def _chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """Splits raw text into 512-token chunks with 64-token overlap."""
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        
        # For token count approximation, we can use character count split (1 token ~ 4 chars)
        # Therefore chunk_size = 512 * 4 = 2048 chars, chunk_overlap = 64 * 4 = 256 chars
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=250,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        splits = splitter.split_text(text)
        chunks = []
        for split in splits:
            # Token count approximation
            token_count = len(split.split())
            chunks.append({
                "content": split,
                "token_count": token_count
            })
            
        return chunks
